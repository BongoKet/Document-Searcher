#!/usr/bin/env python3
"""
Excel Search v5 - Python Search Backend

Communicates with the Electron frontend via newline-delimited JSON
over stdin/stdout. Spawn with `python -u search.py` to ensure
unbuffered output for real-time streaming.

Protocol (input):
  {"type": "search",   "folders": [...], "query": "...",
   "case_sensitive": bool, "search_excel": bool, "search_word": bool,
   "search_pdf": bool, "use_index": bool}
  {"type": "preview",  "file": "...", "location": "..."}
  {"type": "index",    "folders": [...], "extensions": [...]}
  {"type": "index_status", "folders": [...]}
  {"type": "clear_index"}

Protocol (output):
  {"type": "total",      "total": int}
  {"type": "progress",   "file": "...", "scanned": int, "total": int}
  {"type": "result",     "file": "...", "sheet": "...",
                         "cell": "...", "value": "..."}
  {"type": "file_error", "file": "...", "error": "..."}
  {"type": "done",       "scanned": int}
  {"type": "preview_data", "file": "...", "content": ..., "format": "table"|"text"}
  {"type": "index_progress", "file": "...", "indexed": int, "total": int}
  {"type": "index_done", "indexed": int}
  {"type": "index_status", ...}
"""

import sys
import os
import json

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def send(obj: dict) -> None:
    sys.stdout.write(json.dumps(obj, default=str) + "\n")
    sys.stdout.flush()


def col_to_letter(col_idx: int) -> str:
    result = ""
    col_idx += 1
    while col_idx > 0:
        col_idx, rem = divmod(col_idx - 1, 26)
        result = chr(65 + rem) + result
    return result


def cell_ref(row_idx: int, col_idx: int) -> str:
    return f"{col_to_letter(col_idx)}{row_idx + 1}"


def is_match(value, query: str, case_sensitive: bool) -> bool:
    if value is None:
        return False
    val_str = str(value)
    if not val_str.strip():
        return False
    return query in val_str if case_sensitive else query.lower() in val_str.lower()


# ---------------------------------------------------------------------------
# File searchers
# ---------------------------------------------------------------------------

def search_xlsx(filepath: str, query: str, case_sensitive: bool) -> list:
    import openpyxl
    results = []
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if is_match(cell.value, query, case_sensitive):
                        results.append({
                            "sheet": sheet_name,
                            "cell": cell.coordinate,
                            "value": str(cell.value)[:500],
                        })
    finally:
        wb.close()
    return results


def search_xls(filepath: str, query: str, case_sensitive: bool) -> list:
    import xlrd
    results = []
    wb = xlrd.open_workbook(filepath)
    for i in range(wb.nsheets):
        ws = wb.sheet_by_index(i)
        for row_idx in range(ws.nrows):
            for col_idx in range(ws.ncols):
                val = ws.cell_value(row_idx, col_idx)
                if is_match(val, query, case_sensitive):
                    results.append({
                        "sheet": ws.name,
                        "cell": cell_ref(row_idx, col_idx),
                        "value": str(val)[:500],
                    })
    return results


def search_docx(filepath: str, query: str, case_sensitive: bool) -> list:
    from docx import Document
    results = []
    doc = Document(filepath)

    for i, para in enumerate(doc.paragraphs):
        if is_match(para.text, query, case_sensitive):
            results.append({
                "sheet": "Body",
                "cell": f"Paragraph {i + 1}",
                "value": para.text[:500],
            })

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if is_match(cell.text, query, case_sensitive):
                    results.append({
                        "sheet": f"Table {t_idx + 1}",
                        "cell": f"Row {r_idx + 1}, Col {c_idx + 1}",
                        "value": cell.text[:500],
                    })
    return results


def search_pdf(filepath: str, query: str, case_sensitive: bool) -> list:
    import fitz
    results = []
    doc = fitz.open(filepath)
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            for b_idx, block in enumerate(blocks):
                if block[6] != 0:  # skip image blocks
                    continue
                text = block[4]
                if is_match(text, query, case_sensitive):
                    results.append({
                        "sheet": f"Page {page_num + 1}",
                        "cell": f"Block {b_idx + 1}",
                        "value": text.strip()[:500],
                    })
    finally:
        doc.close()
    return results


# ---------------------------------------------------------------------------
# Content extractors (for indexing)
# ---------------------------------------------------------------------------

def extract_xlsx(filepath: str) -> list:
    import openpyxl
    records = []
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    val = cell.value
                    if val is not None and str(val).strip():
                        records.append({
                            "section": sheet_name,
                            "location": cell.coordinate,
                            "text": str(val)[:500],
                        })
    finally:
        wb.close()
    return records


def extract_xls(filepath: str) -> list:
    import xlrd
    records = []
    wb = xlrd.open_workbook(filepath)
    for i in range(wb.nsheets):
        ws = wb.sheet_by_index(i)
        for row_idx in range(ws.nrows):
            for col_idx in range(ws.ncols):
                val = ws.cell_value(row_idx, col_idx)
                if val is not None and str(val).strip():
                    records.append({
                        "section": ws.name,
                        "location": cell_ref(row_idx, col_idx),
                        "text": str(val)[:500],
                    })
    return records


def extract_docx(filepath: str) -> list:
    from docx import Document
    records = []
    doc = Document(filepath)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            records.append({
                "section": "Body",
                "location": f"Paragraph {i + 1}",
                "text": para.text[:500],
            })
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    records.append({
                        "section": f"Table {t_idx + 1}",
                        "location": f"Row {r_idx + 1}, Col {c_idx + 1}",
                        "text": cell.text[:500],
                    })
    return records


def extract_pdf(filepath: str) -> list:
    import fitz
    records = []
    doc = fitz.open(filepath)
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            for b_idx, block in enumerate(blocks):
                if block[6] != 0:
                    continue
                text = block[4].strip()
                if text:
                    records.append({
                        "section": f"Page {page_num + 1}",
                        "location": f"Block {b_idx + 1}",
                        "text": text[:500],
                    })
    finally:
        doc.close()
    return records


# ---------------------------------------------------------------------------
# Preview extractors
# ---------------------------------------------------------------------------

def preview_xlsx(filepath: str, location: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(max_row=100, max_col=26):
                rows.append([
                    str(cell.value) if cell.value is not None else ""
                    for cell in row
                ])
            if any(location == cell_val for row in rows for cell_val in row) or True:
                return {"content": rows, "format": "table"}
    finally:
        wb.close()
    return {"content": [], "format": "table"}


def preview_xls(filepath: str, location: str) -> dict:
    import xlrd
    wb = xlrd.open_workbook(filepath)
    for i in range(wb.nsheets):
        ws = wb.sheet_by_index(i)
        rows = []
        for row_idx in range(min(ws.nrows, 100)):
            row = []
            for col_idx in range(min(ws.ncols, 26)):
                val = ws.cell_value(row_idx, col_idx)
                row.append(str(val) if val else "")
            rows.append(row)
        return {"content": rows, "format": "table"}
    return {"content": [], "format": "table"}


def preview_docx(filepath: str, location: str) -> dict:
    from docx import Document
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]

    para_num = 0
    if location.startswith("Paragraph "):
        try:
            para_num = int(location.split(" ")[1]) - 1
        except (ValueError, IndexError):
            pass

    start = max(0, para_num - 5)
    end = min(len(paragraphs), para_num + 6)
    context = "\n\n".join(paragraphs[start:end])
    return {"content": context, "format": "text"}


def preview_pdf(filepath: str, location: str) -> dict:
    import fitz
    doc = fitz.open(filepath)
    try:
        page_num = 0
        if "Page " in location:
            try:
                page_num = int(location.split("Page ")[1].split(",")[0].split(" ")[0]) - 1
            except (ValueError, IndexError):
                pass

        if 0 <= page_num < len(doc):
            page = doc[page_num]
            text = page.get_text("text")
            return {"content": text, "format": "text"}
    finally:
        doc.close()
    return {"content": "", "format": "text"}


# ---------------------------------------------------------------------------
# Registries
# ---------------------------------------------------------------------------

SEARCHERS = {
    ".xlsx": search_xlsx,
    ".xls":  search_xls,
    ".docx": search_docx,
    ".pdf":  search_pdf,
}

EXTRACTORS = {
    ".xlsx": extract_xlsx,
    ".xls":  extract_xls,
    ".docx": extract_docx,
    ".pdf":  extract_pdf,
}

PREVIEWERS = {
    ".xlsx": preview_xlsx,
    ".xls":  preview_xls,
    ".docx": preview_docx,
    ".pdf":  preview_pdf,
}


# ---------------------------------------------------------------------------
# Main search loop
# ---------------------------------------------------------------------------

def gather_extensions(search_excel: bool, search_word: bool, search_pdf_flag: bool) -> list:
    exts = []
    if search_excel:
        exts += [".xlsx", ".xls"]
    if search_word:
        exts += [".docx"]
    if search_pdf_flag:
        exts += [".pdf"]
    return exts


def count_files(folders: list, extensions: list) -> int:
    total = 0
    for folder in folders:
        for root, dirs, files in os.walk(folder):
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            total += sum(1 for f in files if os.path.splitext(f)[1].lower() in extensions)
    return total


def collect_files(folders: list, extensions: list):
    """Yield (filepath, ext) for all matching files across folders."""
    for folder in folders:
        for root, dirs, files in os.walk(folder):
            dirs[:] = sorted(d for d in dirs if not d.startswith("."))
            for filename in sorted(files):
                ext = os.path.splitext(filename)[1].lower()
                if ext in extensions:
                    yield os.path.join(root, filename), ext


def run_search(folders: list, query: str, case_sensitive: bool,
               search_excel: bool, search_word: bool,
               search_pdf_flag: bool, use_index: bool = False) -> None:
    extensions = gather_extensions(search_excel, search_word, search_pdf_flag)
    if not extensions:
        send({"type": "done", "scanned": 0})
        return

    if use_index:
        try:
            from indexer import SearchIndex
            idx = SearchIndex()
            indexed_results = idx.search(query, case_sensitive, folders)
            send({"type": "total", "total": len(indexed_results)})
            for i, r in enumerate(indexed_results):
                send({
                    "type": "result",
                    "file": r["file"],
                    "sheet": r["section"],
                    "cell": r["location"],
                    "value": r["value"],
                })
                if (i + 1) % 50 == 0:
                    send({"type": "progress", "file": r["file"],
                          "scanned": i + 1, "total": len(indexed_results)})
            send({"type": "done", "scanned": len(indexed_results)})
            idx.close()
            return
        except Exception as exc:
            send({"type": "file_error", "file": "index", "error": f"Index search failed, falling back to live scan: {exc}"})

    total = count_files(folders, extensions)
    send({"type": "total", "total": total})

    scanned = 0
    for filepath, ext in collect_files(folders, extensions):
        scanned += 1
        send({"type": "progress", "file": filepath, "scanned": scanned, "total": total})
        try:
            searcher = SEARCHERS[ext]
            for result in searcher(filepath, query, case_sensitive):
                send({"type": "result", "file": filepath, **result})
        except Exception as exc:
            send({"type": "file_error", "file": filepath, "error": str(exc)})

    send({"type": "done", "scanned": scanned})


# ---------------------------------------------------------------------------
# Indexing
# ---------------------------------------------------------------------------

def run_index(folders: list, extensions: list) -> None:
    from indexer import SearchIndex
    idx = SearchIndex()

    all_files = list(collect_files(folders, extensions))
    total = len(all_files)
    send({"type": "index_progress", "file": "", "indexed": 0, "total": total})

    indexed = 0
    for filepath, ext in all_files:
        if idx.is_fresh(filepath):
            indexed += 1
            continue
        try:
            extractor = EXTRACTORS.get(ext)
            if extractor:
                records = extractor(filepath)
                idx.index_file(filepath, records)
        except Exception as exc:
            send({"type": "file_error", "file": filepath, "error": str(exc)})

        indexed += 1
        if indexed % 10 == 0 or indexed == total:
            send({"type": "index_progress", "file": filepath,
                  "indexed": indexed, "total": total})

    idx.close()
    send({"type": "index_done", "indexed": indexed})


def run_index_status(folders: list) -> None:
    try:
        from indexer import SearchIndex
        idx = SearchIndex()
        status = idx.status(folders if folders else None)
        idx.close()
        send({"type": "index_status", **status})
    except Exception as exc:
        send({"type": "index_status", "total_files": 0, "total_content": 0,
              "error": str(exc)})


def run_clear_index() -> None:
    try:
        from indexer import SearchIndex
        idx = SearchIndex()
        idx.clear()
        idx.close()
        send({"type": "index_cleared"})
    except Exception as exc:
        send({"type": "error", "message": f"Failed to clear index: {exc}"})


# ---------------------------------------------------------------------------
# Preview
# ---------------------------------------------------------------------------

def run_preview(filepath: str, location: str) -> None:
    ext = os.path.splitext(filepath)[1].lower()
    previewer = PREVIEWERS.get(ext)
    if not previewer:
        send({"type": "preview_data", "file": filepath,
              "content": f"Preview not available for {ext} files", "format": "text"})
        return
    try:
        result = previewer(filepath, location)
        send({"type": "preview_data", "file": filepath, **result})
    except Exception as exc:
        send({"type": "preview_data", "file": filepath,
              "content": f"Preview error: {exc}", "format": "text"})


# ---------------------------------------------------------------------------
# Main loop
# ---------------------------------------------------------------------------

def main() -> None:
    for raw_line in sys.stdin:
        line = raw_line.strip()
        if not line:
            continue
        try:
            cmd = json.loads(line)
        except json.JSONDecodeError:
            continue

        cmd_type = cmd.get("type")

        if cmd_type == "search":
            folders = cmd.get("folders") or []
            if not folders and cmd.get("folder"):
                folders = [cmd["folder"]]
            run_search(
                folders=folders,
                query=cmd.get("query", ""),
                case_sensitive=cmd.get("case_sensitive", False),
                search_excel=cmd.get("search_excel", True),
                search_word=cmd.get("search_word", True),
                search_pdf_flag=cmd.get("search_pdf", True),
                use_index=cmd.get("use_index", False),
            )

        elif cmd_type == "preview":
            run_preview(
                filepath=cmd.get("file", ""),
                location=cmd.get("location", ""),
            )

        elif cmd_type == "index":
            folders = cmd.get("folders", [])
            extensions = cmd.get("extensions", [".xlsx", ".xls", ".docx", ".pdf"])
            run_index(folders, extensions)

        elif cmd_type == "index_status":
            run_index_status(cmd.get("folders", []))

        elif cmd_type == "clear_index":
            run_clear_index()


if __name__ == "__main__":
    main()
